import csv
import os

import fitz  # pymupdf

from app.config import settings
from app.tasks.celery_app import celery_app
from app.tasks.nodes.base import mark_step_running, mark_step_done, mark_step_failed, mark_run_failed, step_log, raise_if_cancelled

MAX_TEXT_CHARS = 50_000  # ~12 K tokens — keeps any single document within model limits


# ── Text extractors ───────────────────────────────────────────────────────────

def _extract_docx_text(file_path: str) -> str:
    from docx import Document  # python-docx
    doc = Document(file_path)
    paras = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also pull table cells
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                paras.append(" | ".join(cells))
    return "\n".join(paras)


def _extract_xlsx_text(file_path: str) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
    lines = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        lines.append(f"[Sheet: {sheet_name}]")
        for row in ws.iter_rows(values_only=True):
            row_str = "\t".join("" if v is None else str(v) for v in row)
            if row_str.strip():
                lines.append(row_str)
    wb.close()
    return "\n".join(lines)


def _extract_csv_text(file_path: str) -> str:
    with open(file_path, newline="", encoding="utf-8-sig", errors="replace") as f:
        reader = csv.reader(f)
        return "\n".join("\t".join(row) for row in reader)


_TEXT_EXTRACTORS: dict = {
    ".docx": _extract_docx_text,
    ".xlsx": _extract_xlsx_text,
    ".xls":  _extract_xlsx_text,
    ".csv":  _extract_csv_text,
}


# ── Image converter ───────────────────────────────────────────────────────────

def _convert_file_to_images(file_path: str, out_dir: str, rel_dir: str, scale: float) -> list[str]:
    """Convert a single file (PDF or image) to page PNGs; returns relative paths."""
    os.makedirs(out_dir, exist_ok=True)
    fitz_doc = fitz.open(file_path)
    image_paths = []
    for page_num, page in enumerate(fitz_doc):
        mat = fitz.Matrix(scale, scale)
        pix = page.get_pixmap(matrix=mat)
        img_file = f"page_{page_num + 1:04d}.png"
        img_path = os.path.join(out_dir, img_file)
        pix.save(img_path)
        image_paths.append(f"{rel_dir}/{img_file}")
    fitz_doc.close()
    return image_paths


# ── Task ──────────────────────────────────────────────────────────────────────

@celery_app.task(name="nodes.pdf_to_images", bind=True)
def pdf_to_images_task(self, input_data: dict, run_id: int, step_id: int, node_config: dict | None = None) -> dict:
    mark_step_running(step_id)
    try:
        raise_if_cancelled(run_id)
        cfg = node_config or {}
        scale = float(cfg.get("scale", 2.0))
        step_log(step_id, f"Scale: {scale}x")

        documents_list: list[dict] | None = input_data.get("documents")

        if documents_list:
            # ── Multi-doc mode ────────────────────────────────────────────────
            result_docs = []
            for doc_info in documents_list:
                doc_id = doc_info["id"]
                file_path = doc_info["file_path"]
                filename = doc_info.get("filename", f"doc_{doc_id}")
                ext = os.path.splitext(file_path)[1].lower()

                step_log(step_id, f"Processing: {filename}")

                if ext in _TEXT_EXTRACTORS:
                    # ── Text format — extract content ──────────────────────
                    try:
                        text = _TEXT_EXTRACTORS[ext](file_path)
                        if len(text) > MAX_TEXT_CHARS:
                            text = text[:MAX_TEXT_CHARS] + f"\n\n[... truncated at {MAX_TEXT_CHARS} chars]"
                        step_log(step_id, f"  {filename}: extracted {len(text):,} chars of text")
                        result_docs.append({
                            "id": doc_id,
                            "filename": filename,
                            "image_paths": [],
                            "text_content": text,
                        })
                    except Exception as exc:
                        step_log(step_id, f"  {filename}: text extraction failed — {exc}")
                        result_docs.append({
                            "id": doc_id,
                            "filename": filename,
                            "image_paths": [],
                            "text_content": f"[Could not extract text from {filename}: {exc}]",
                        })
                else:
                    # ── Image/PDF format — render to pages ─────────────────
                    rel_dir = f"run_{run_id}_doc_{doc_id}_pages"
                    out_dir = os.path.join(settings.storage_path, rel_dir)

                    try:
                        image_paths = _convert_file_to_images(file_path, out_dir, rel_dir, scale)
                        step_log(step_id, f"  {filename}: {len(image_paths)} page(s)")
                    except Exception as exc:
                        step_log(step_id, f"  {filename}: conversion failed — {exc}")
                        image_paths = []

                    result_docs.append({
                        "id": doc_id,
                        "filename": filename,
                        "image_paths": image_paths,
                    })

            image_docs = [d for d in result_docs if d.get("image_paths")]
            text_docs  = [d for d in result_docs if d.get("text_content")]
            total_pages = sum(len(d["image_paths"]) for d in image_docs)
            step_log(step_id, (
                f"Total: {len(result_docs)} doc(s) — "
                f"{len(image_docs)} visual ({total_pages} page(s)), "
                f"{len(text_docs)} text"
            ))

            all_image_paths = [p for d in result_docs for p in d.get("image_paths", [])]

            output = {
                "documents": result_docs,
                "document_id": input_data.get("document_id"),
                "image_paths": all_image_paths,
                "page_count": total_pages,
            }

        else:
            # ── Single-doc fallback (workflow-editor chains, email_input, etc.) ──
            file_path = input_data["file_path"]
            document_id = input_data["document_id"]
            ext = os.path.splitext(file_path)[1].lower()

            step_log(step_id, f"File: {os.path.basename(file_path)}")

            if ext in _TEXT_EXTRACTORS:
                # Text format — extract and pass through as text_content
                try:
                    text = _TEXT_EXTRACTORS[ext](file_path)
                    if len(text) > MAX_TEXT_CHARS:
                        text = text[:MAX_TEXT_CHARS] + f"\n\n[... truncated at {MAX_TEXT_CHARS} chars]"
                    step_log(step_id, f"Extracted {len(text):,} chars of text")
                except Exception as exc:
                    step_log(step_id, f"Text extraction failed — {exc}")
                    text = f"[Could not extract text: {exc}]"

                output = {
                    "document_id": document_id,
                    "file_path": file_path,
                    "mime_type": input_data.get("mime_type"),
                    "image_paths": [],
                    "page_count": 0,
                    "text_content": text,
                }
            else:
                rel_dir = f"run_{run_id}_doc_{document_id}_pages"
                out_dir = os.path.join(settings.storage_path, rel_dir)
                os.makedirs(out_dir, exist_ok=True)

                fitz_doc = fitz.open(file_path)
                page_count = len(fitz_doc)
                step_log(step_id, f"Converting: {page_count} page(s)")

                image_paths = []
                for page_num, page in enumerate(fitz_doc):
                    mat = fitz.Matrix(scale, scale)
                    pix = page.get_pixmap(matrix=mat)
                    img_file = f"page_{page_num + 1:04d}.png"
                    img_path = os.path.join(out_dir, img_file)
                    pix.save(img_path)
                    image_paths.append(f"{rel_dir}/{img_file}")
                fitz_doc.close()

                step_log(step_id, f"Saved {len(image_paths)} image(s) to {rel_dir}/")

                output = {
                    "document_id": document_id,
                    "image_paths": image_paths,
                    "page_count": len(image_paths),
                }

        mark_step_done(step_id, output)
        return output

    except Exception as exc:
        mark_step_failed(step_id, str(exc))
        mark_run_failed(run_id, str(exc))
        raise
