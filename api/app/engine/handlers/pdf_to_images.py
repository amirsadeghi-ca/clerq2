"""`pdf_to_images` node — render PDFs/images to page PNGs, extract text from
Word/Excel/CSV. Multi-doc only (the single-doc dual path is gone); files go
through ctx.storage."""
from __future__ import annotations

import csv
import os

import fitz  # pymupdf

from app.engine.context import Output, StepContext

MAX_TEXT_CHARS = 50_000  # ~12K tokens — keeps any single document within model limits


def _extract_docx_text(path: str) -> str:
    from docx import Document
    doc = Document(path)
    paras = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                paras.append(" | ".join(cells))
    return "\n".join(paras)


def _extract_xlsx_text(path: str) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    lines = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        lines.append(f"[Sheet: {sheet_name}]")
        for row in ws.iter_rows(values_only=True):
            row_str = "\t".join("" if v is None else str(v) for v in row)
            if row_str.strip():
                lines.append(row_str)
    wb.close()
    return "\n".join(lines)


def _extract_csv_text(path: str) -> str:
    with open(path, newline="", encoding="utf-8-sig", errors="replace") as f:
        return "\n".join("\t".join(row) for row in csv.reader(f))


_TEXT_EXTRACTORS = {
    ".docx": _extract_docx_text,
    ".xlsx": _extract_xlsx_text,
    ".xls": _extract_xlsx_text,
    ".csv": _extract_csv_text,
}


def pdf_to_images(ctx: StepContext) -> Output:
    cfg = ctx.config or {}
    scale = float(cfg.get("scale", 2.0))
    ctx.log(f"Scale: {scale}x")

    result_docs: list[dict] = []
    for d in ctx.documents():
        ctx.check_cancelled()
        doc_id = d.get("id")
        file_path = d.get("file_path") or ""
        name = d.get("filename", f"doc_{doc_id}")
        ext = ("." + file_path.rsplit(".", 1)[-1].lower()) if "." in file_path else ""
        abs_in = file_path if os.path.isabs(file_path) else ctx.storage.abspath(file_path)

        if ext in _TEXT_EXTRACTORS:
            try:
                text = _TEXT_EXTRACTORS[ext](abs_in)
                if len(text) > MAX_TEXT_CHARS:
                    text = text[:MAX_TEXT_CHARS] + f"\n\n[... truncated at {MAX_TEXT_CHARS} chars]"
                ctx.log(f"  {name}: extracted {len(text):,} chars of text")
                result_docs.append({"id": doc_id, "filename": name, "image_paths": [], "text_content": text})
            except Exception as exc:  # noqa: BLE001
                ctx.log(f"  {name}: text extraction failed — {exc}")
                result_docs.append({"id": doc_id, "filename": name, "image_paths": [],
                                    "text_content": f"[Could not extract text from {name}: {exc}]"})
        else:
            rel_dir = f"run_{ctx.run_id}_doc_{doc_id}_pages"
            paths: list[str] = []
            try:
                fdoc = fitz.open(abs_in)
                for i, page in enumerate(fdoc):
                    pix = page.get_pixmap(matrix=fitz.Matrix(scale, scale))
                    rel = f"{rel_dir}/page_{i + 1:04d}.png"
                    ctx.storage.save(rel, pix.tobytes("png"))
                    paths.append(rel)
                fdoc.close()
                ctx.log(f"  {name}: {len(paths)} page(s)")
            except Exception as exc:  # noqa: BLE001
                ctx.log(f"  {name}: conversion failed — {exc}")
            result_docs.append({"id": doc_id, "filename": name, "image_paths": paths})

    all_paths = [p for d in result_docs for p in d.get("image_paths", [])]
    return Output({
        "documents": result_docs,
        "image_paths": all_paths,
        "page_count": len(all_paths),
        "document_id": ctx.primary_input().get("document_id"),
    })
